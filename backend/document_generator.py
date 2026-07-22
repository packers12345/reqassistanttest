from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from io import BytesIO


def generate_feedback_document(
    session_id: str,
    analysis: Dict,
    feedback: Dict
) -> bytes:
    """
    Generate comprehensive Word document (returned as in-memory bytes) with:
    - All requirements
    - AI suggestions
    - User decisions
    - Final text

    Nothing is written to disk — the caller holds the bytes in memory.
    """
    doc = Document()

    # Set up default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('INCOSE Requirements Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Session ID: {session_id}")
    doc.add_paragraph(f"Analysis Date: {analysis.get('timestamp', 'N/A')}")
    doc.add_paragraph(f"Total Requirements: {len(analysis.get('requirements', []))}")

    stats = feedback.get('summary_statistics', {})
    if stats:
        doc.add_paragraph(f"Total Violations Found: {stats.get('total_violations', 0)}")
        actions = stats.get('actions', {})
        doc.add_paragraph(
            f"Actions: {actions.get('accept', 0)} Accepted, "
            f"{actions.get('reject', 0)} Rejected, "
            f"{actions.get('modify', 0)} Modified"
        )

    doc.add_page_break()

    # Process each requirement
    for req_feedback in feedback.get('requirement_feedback', []):
        req_id = req_feedback['req_id']

        # Find corresponding analysis
        req_analysis = next(
            (r for r in analysis.get('requirements', []) if r['req_id'] == req_id),
            None
        )

        if not req_analysis:
            continue

        # Requirement header
        doc.add_heading(f'Requirement {req_id}', 1)

        # Original text
        doc.add_heading('ORIGINAL TEXT:', 3)
        p = doc.add_paragraph(req_analysis['original_text'])
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)

        # Violations section
        violations_found = len(req_feedback.get('violation_feedback', []))
        doc.add_heading(f'VIOLATIONS FOUND: {violations_found}', 3)

        # The analysis stores per-criterion results under 'criteria_evaluations'
        # keyed by criterion_id; feedback carries that id as 'rule_id'.
        evaluations = {
            ev.get('criterion_id'): ev
            for ev in req_analysis.get('criteria_evaluations', [])
        }

        for viol_feedback in req_feedback.get('violation_feedback', []):
            rule_id = viol_feedback.get('rule_id', '')
            violation = evaluations.get(rule_id, {})

            # Rule header
            doc.add_heading(
                f"Violation: {rule_id} - "
                f"{viol_feedback.get('criterion_name') or violation.get('criterion_name', '')}",
                4
            )

            # Violation details table
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'

            # Explanation
            table.rows[0].cells[0].text = 'Problem:'
            table.rows[0].cells[1].text = violation.get('explanation', '')

            # Affected text
            table.rows[1].cells[0].text = 'Affected Text:'
            cell = table.rows[1].cells[1]
            cell.text = violation.get('affected_text') or ''
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)

            # AI suggestion
            table.rows[2].cells[0].text = 'AI Suggestion:'
            cell = table.rows[2].cells[1]
            cell.text = (
                viol_feedback.get('ai_suggestion')
                or violation.get('suggested_replacement')
                or ''
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 128, 0)

            # User decision
            table.rows[3].cells[0].text = 'User Decision:'
            action = viol_feedback.get('user_action', '').upper()

            if action == 'ACCEPT':
                color = RGBColor(0, 128, 0)
            elif action == 'REJECT':
                color = RGBColor(255, 0, 0)
            else:
                color = RGBColor(255, 165, 0)

            cell = table.rows[3].cells[1]
            cell.text = action
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
                    run.font.bold = True

            # Final text for this violation
            table.rows[4].cells[0].text = 'Final Text:'
            cell = table.rows[4].cells[1]
            cell.text = viol_feedback.get('user_text', '')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

            # User notes
            if viol_feedback.get('notes'):
                doc.add_paragraph(
                    f"Notes: {viol_feedback['notes']}",
                    style='Intense Quote'
                )

            doc.add_paragraph()

        # Reviewer suggestions on criteria the AI marked satisfied
        suggestions = req_feedback.get('suggestion_feedback', [])
        if suggestions:
            doc.add_heading(f'REVIEWER SUGGESTIONS: {len(suggestions)}', 3)
            for sug in suggestions:
                doc.add_heading(
                    f"Suggestion: {sug.get('rule_id', '')} - {sug.get('criterion_name', '')}",
                    4
                )
                p = doc.add_paragraph()
                run = p.add_run(sug.get('user_text', ''))
                run.font.bold = True
                if sug.get('notes'):
                    doc.add_paragraph(f"Notes: {sug['notes']}", style='Intense Quote')
            doc.add_paragraph()

        # Final requirement text
        doc.add_heading('FINAL REQUIREMENT:', 3)
        p = doc.add_paragraph()
        run = p.add_run(req_feedback.get('final_text', ''))
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 255)

        if req_feedback.get('overall_notes'):
            doc.add_paragraph(
                f"Overall Notes: {req_feedback['overall_notes']}",
                style='Intense Quote'
            )

        doc.add_paragraph()

    # Summary statistics
    if stats:
        doc.add_page_break()
        doc.add_heading('Summary Statistics', 1)

        doc.add_paragraph(f"Total Requirements Analyzed: {stats.get('total_requirements', 0)}")
        doc.add_paragraph(f"Total Violations Found: {stats.get('total_violations', 0)}")

        actions = stats.get('actions', {})
        doc.add_paragraph(
            f"User Actions: {actions.get('accept', 0)} Accepted, "
            f"{actions.get('reject', 0)} Rejected, "
            f"{actions.get('modify', 0)} Modified"
        )

        rate_by_rule = stats.get('acceptance_rate_by_rule', {})
        if rate_by_rule:
            doc.add_heading('Acceptance Rate by INCOSE Rule:', 3)

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Rule ID'
            hdr_cells[1].text = 'Accepted'
            hdr_cells[2].text = 'Rejected'
            hdr_cells[3].text = 'Modified'
            hdr_cells[4].text = 'Rate'

            for rule_id, rule_stats in rate_by_rule.items():
                row_cells = table.add_row().cells
                row_cells[0].text = rule_id
                row_cells[1].text = str(rule_stats.get('accept', 0))
                row_cells[2].text = str(rule_stats.get('reject', 0))
                row_cells[3].text = str(rule_stats.get('modify', 0))
                row_cells[4].text = f"{rule_stats.get('rate', 0):.1%}"

    # Serialize document to in-memory bytes (no disk write)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
